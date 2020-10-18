#! python3
# invitations.py - generates custom invitations addressed to the guests listed in guests.txt

import docx,os,re

# get the list of guests 
guest_file = open('guests.txt','r')
guest_list = guest_file.readlines()

# open the invitations document
invitations_docx = docx.Document('invitations.docx')

# remove the newlines from each guest's name
for guest in guest_list:
    guest = re.sub('\n','',guest)
    
    # create an invitation for each guest_file
    print(f"Writing {guest}'s invitation...")
    pleasure_line = invitations_docx.add_paragraph('It would be a pleasure to have the company of','Heading1')
    guest_line = invitations_docx.add_paragraph(guest, 'Heading3')
    address_line = invitations_docx.add_paragraph('at 11010 Memory Lane on the Evening of','Heading1')
    date_line = invitations_docx.add_paragraph('April 1st','Heading2')
    time_line = invitations_docx.add_paragraph('at 7 o’clock','Heading1')

    #insert a page break for the next invite
    time_line.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    print(f"Wrote {guest}'s invitation.")

# save the invitations document
invitations_docx.save('invitations.docx')
print('Saved invitations.')
